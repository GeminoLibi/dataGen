"""
File Generator - Creates actual files in various formats (PDF, DOCX, XLSX, etc.)
"""

import os
import random
from datetime import datetime, timedelta
from typing import Dict, List, Optional
from faker import Faker

fake = Faker()


class FileGenerator:
    """Generates actual files in various formats."""
    
    def __init__(self, output_dir: str):
        self.output_dir = output_dir
        self.file_counter = 1
    
    def generate_pdf(self, content: str, filename: Optional[str] = None) -> str:
        """Generate a PDF file."""
        try:
            from reportlab.lib.pagesizes import letter
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.styles import getSampleStyleSheet
            
            if not filename:
                filename = f"document_{self.file_counter:03d}.pdf"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = SimpleDocTemplate(filepath, pagesize=letter)
            styles = getSampleStyleSheet()
            story = []
            
            # Split content into paragraphs
            for line in content.split('\n'):
                if line.strip():
                    story.append(Paragraph(line.strip(), styles['Normal']))
                else:
                    story.append(Spacer(1, 12))
            
            doc.build(story)
            self.file_counter += 1
            return filepath
        except ImportError:
            # Fallback to text file if reportlab not available
            return self.generate_txt(content, filename.replace('.pdf', '.txt') if filename else None)
    
    def generate_docx(self, content: str, filename: Optional[str] = None) -> str:
        """Generate a DOCX file."""
        try:
            from docx import Document
            
            if not filename:
                filename = f"document_{self.file_counter:03d}.docx"
            filepath = os.path.join(self.output_dir, filename)
            
            doc = Document()
            
            # Split content into paragraphs
            for line in content.split('\n'):
                if line.strip():
                    doc.add_paragraph(line.strip())
                else:
                    doc.add_paragraph()
            
            doc.save(filepath)
            self.file_counter += 1
            return filepath
        except ImportError:
            # Fallback to text file
            return self.generate_txt(content, filename.replace('.docx', '.txt') if filename else None)
    
    def generate_xlsx(self, data: List[List[str]], headers: List[str], filename: Optional[str] = None) -> str:
        """Generate an XLSX file with actual data."""
        try:
            from openpyxl import Workbook
            
            if not filename:
                filename = f"data_{self.file_counter:03d}.xlsx"
            filepath = os.path.join(self.output_dir, filename)
            
            wb = Workbook()
            ws = wb.active
            
            # Add headers
            ws.append(headers)
            
            # Add data rows
            for row in data:
                ws.append(row)
            
            wb.save(filepath)
            self.file_counter += 1
            return filepath
        except ImportError:
            # Fallback to CSV
            return self.generate_csv(data, headers, filename.replace('.xlsx', '.csv') if filename else None)
    
    def generate_csv(self, data: List[List[str]], headers: List[str], filename: Optional[str] = None) -> str:
        """Generate a CSV file."""
        import csv
        
        if not filename:
            filename = f"data_{self.file_counter:03d}.csv"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            writer.writerows(data)
        
        self.file_counter += 1
        return filepath
    
    def generate_txt(self, content: str, filename: Optional[str] = None) -> str:
        """Generate a text file."""
        if not filename:
            filename = f"document_{self.file_counter:03d}.txt"
        filepath = os.path.join(self.output_dir, filename)
        
        with open(filepath, 'w', encoding='utf-8') as f:
            f.write(content)
        
        self.file_counter += 1
        return filepath
    
    def generate_financial_xlsx(self, num_transactions: int = 50) -> str:
        """Generate financial records as XLSX."""
        headers = ['Date', 'Description', 'Amount', 'Account', 'Category', 'Balance']
        data = []
        
        balance = random.randint(1000, 50000)
        for i in range(num_transactions):
            date = (datetime.now() - timedelta(days=random.randint(0, 90))).strftime('%Y-%m-%d')
            descriptions = [
                f"ATM WITHDRAWAL - {fake.city()}",
                f"PURCHASE - {fake.company()}",
                f"TRANSFER TO {fake.name()}",
                f"DEPOSIT - CHECK #{random.randint(1000, 9999)}",
                f"ONLINE PAYMENT - {fake.company()}",
                f"CASH DEPOSIT",
                f"FEE - {random.choice(['MONTHLY', 'OVERDRAFT', 'WIRE'])}"
            ]
            amount = random.randint(-5000, 3000)
            balance += amount
            account = f"****{random.randint(1000, 9999)}"
            category = random.choice(['ATM', 'PURCHASE', 'TRANSFER', 'DEPOSIT', 'FEE'])
            
            data.append([date, random.choice(descriptions), f"{amount:.2f}", account, category, f"{balance:.2f}"])
        
        return self.generate_xlsx(data, headers, f"financial_records_{self.file_counter:03d}.xlsx")
    
    def generate_evidence_log_xlsx(self, evidence_items: List[Dict]) -> str:
        """Generate evidence log as XLSX."""
        headers = ['Evidence ID', 'Type', 'Description', 'Location Found', 'Date Collected', 'Collected By', 'Status']
        data = []
        
        evidence_types = ['Physical', 'Digital', 'Document', 'Biological', 'Firearm']
        locations = [fake.address() for _ in range(5)]
        officers = [fake.name() for _ in range(3)]
        statuses = ['In Storage', 'At Lab', 'Returned']
        
        for item in evidence_items:
            evid_id = item.get('id', f"EVID-{random.randint(1000, 9999)}")
            evid_type = random.choice(evidence_types)
            description = item.get('description', 'Evidence recovered from scene')
            location = random.choice(locations)
            date = (datetime.now() - timedelta(days=random.randint(0, 30))).strftime('%Y-%m-%d')
            officer = random.choice(officers)
            status = random.choice(statuses)
            
            data.append([evid_id, evid_type, description, location, date, officer, status])
        
        return self.generate_xlsx(data, headers, f"evidence_log_{self.file_counter:03d}.xlsx")
    
    def generate_phone_records_xlsx(self, num_calls: int = 100) -> str:
        """Generate phone records as XLSX."""
        headers = ['Date', 'Time', 'Duration', 'From Number', 'To Number', 'Call Type', 'Location']
        data = []
        
        for i in range(num_calls):
            date = (datetime.now() - timedelta(days=random.randint(0, 30))).strftime('%Y-%m-%d')
            time = f"{random.randint(0, 23):02d}:{random.randint(0, 59):02d}"
            duration = random.randint(10, 3600)
            from_num = f"{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}"
            to_num = f"{random.randint(200, 999)}-{random.randint(200, 999)}-{random.randint(1000, 9999)}"
            call_type = random.choice(['Voice', 'Text', 'Data'])
            location = f"{fake.city()}, {fake.state_abbr()}"
            
            data.append([date, time, str(duration), from_num, to_num, call_type, location])
        
        return self.generate_xlsx(data, headers, f"phone_records_{self.file_counter:03d}.xlsx")
    
    def generate_incident_report_pdf(self, content: str) -> str:
        """Generate incident report as PDF."""
        return self.generate_pdf(content, f"incident_report_{self.file_counter:03d}.pdf")
    
    def generate_memo_docx(self, content: str) -> str:
        """Generate memo as DOCX."""
        return self.generate_docx(content, f"memo_{self.file_counter:03d}.docx")
    
    def generate_ransom_note_txt(self, content: str) -> str:
        """Generate ransom note as TXT."""
        return self.generate_txt(content, f"ransom_note_{self.file_counter:03d}.txt")

