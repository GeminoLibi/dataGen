#!/usr/bin/env python3
"""
Convert all .docx files in the evals directory to .txt files.
"""

import os
from pathlib import Path
from rich.console import Console

console = Console()

try:
    import docx2txt
    HAS_DOCX2TXT = True
except ImportError:
    HAS_DOCX2TXT = False

try:
    from docx import Document
    HAS_PYTHON_DOCX = True
except ImportError:
    HAS_PYTHON_DOCX = False

def convert_docx_to_txt_docx2txt(docx_path: Path, txt_path: Path) -> bool:
    """Convert using docx2txt library."""
    try:
        text = docx2txt.process(str(docx_path))
        txt_path.write_text(text, encoding='utf-8')
        return True
    except Exception as e:
        console.print(f"[red]Error converting {docx_path.name}: {e}[/red]")
        return False

def convert_docx_to_txt_python_docx(docx_path: Path, txt_path: Path) -> bool:
    """Convert using python-docx library."""
    try:
        doc = Document(str(docx_path))
        text_parts = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(" | ".join(row_text))
        
        full_text = "\n".join(text_parts)
        txt_path.write_text(full_text, encoding='utf-8')
        return True
    except Exception as e:
        console.print(f"[red]Error converting {docx_path.name}: {e}[/red]")
        return False

def main():
    """Convert all .docx files in the evals directory to .txt."""
    evals_dir = Path(__file__).parent
    
    # Check for available libraries
    if not HAS_DOCX2TXT and not HAS_PYTHON_DOCX:
        console.print("[red]No .docx conversion library found![/red]")
        console.print("[yellow]Installing python-docx...[/yellow]")
        import subprocess
        subprocess.run([sys.executable, "-m", "pip", "install", "python-docx"], check=False)
        console.print("[green]Please run this script again.[/green]")
        return
    
    # Find all .docx files
    docx_files = list(evals_dir.glob("*.docx"))
    
    if not docx_files:
        console.print("[yellow]No .docx files found in evals directory.[/yellow]")
        return
    
    console.print(f"[cyan]Found {len(docx_files)} .docx files to convert[/cyan]")
    
    converted = 0
    failed = 0
    
    for docx_path in docx_files:
        txt_path = docx_path.with_suffix('.txt')
        
        # Skip if .txt already exists
        if txt_path.exists():
            console.print(f"[yellow]Skipping {docx_path.name} (already converted)[/yellow]")
            continue
        
        console.print(f"[cyan]Converting: {docx_path.name}[/cyan]")
        
        # Try docx2txt first (better formatting), then python-docx
        success = False
        if HAS_DOCX2TXT:
            success = convert_docx_to_txt_docx2txt(docx_path, txt_path)
        
        if not success and HAS_PYTHON_DOCX:
            success = convert_docx_to_txt_python_docx(docx_path, txt_path)
        
        if success:
            converted += 1
            console.print(f"[green]✓ Converted: {txt_path.name}[/green]")
        else:
            failed += 1
            console.print(f"[red]✗ Failed: {docx_path.name}[/red]")
    
    console.print("")
    console.print(f"[green]Conversion complete![/green]")
    console.print(f"  Converted: {converted}")
    console.print(f"  Failed: {failed}")
    console.print(f"  Total: {len(docx_files)}")

if __name__ == "__main__":
    import sys
    main()

